"""
dont use any vector db. upload the doc and pass it as context
"""
import streamlit as st
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_ollama import OllamaEmbeddings
from langchain_ollama.llms import OllamaLLM
from langchain_docling import DoclingLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
import pandas as pd
from text_to_image_v2 import generate_infographic
import os
from docx import Document

# user_query ="""
# - participant asked to participate in the clinical research because?
# - what is the goal of the study?
# - what is the drug under study?
# - explain how drug works?
# - explain side effects and impacts w.r.to patient health?
# - what are the risks involved?
# - who is the sponsor for the study?
# - what is the medicalcare provided?
# - what is the trial period? (should be number of calendar days)
# - (optional) do you have any statistics for the number of participants participating?
# """
user_query =[
"what is that one important reason one can participate this clinical trial?",
"what is the goal of the study?",
"what is the drug under study?",
"explain how drug works?",
"explain side effects and impacts w.r.to patient health?",
"what are the risks involved?",
"who is the sponsor for the study?",
"what is the medicalcare provided?",
"what is the trial period? (should be number of calendar days)",
"(optional) do you have any statistics for the number of participants participating?",
]

PDF_STORAGE_PATH = './document_store/pdfs/'
EMBEDDING_MODEL = OllamaEmbeddings(model="deepseek-r1:1.5b")
DOCUMENT_VECTOR_DB = InMemoryVectorStore(EMBEDDING_MODEL)
# LANGUAGE_MODEL = OllamaLLM(model="deepseek-r1:1.5b", params={"temperature": 0, "seed": 42, "top_k": 1})
LANGUAGE_MODEL = OllamaLLM(model="llama2:13b", params={"temperature": 0.4, "seed": 42, "top_k": 1})

def extract_text_and_images_and_tables(file_path, image_storage_path='./extracted_images/'):
    document = Document(file_path)
    full_text = []
    tables_text = []
    
    # Ensure the image storage path exists
    if not os.path.exists(image_storage_path):
        os.makedirs(image_storage_path)
    
    # Extract text from paragraphs
    for para in document.paragraphs:
        full_text.append(para.text)
    
    # Extract tables
    for table in document.tables:
        table_text = ""
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            table_text += row_text + "\n"
        tables_text.append(table_text.strip())  # Add table text to tables_text list
    
    # Extract images (if any)
    for rel in document.part.rels.values():
        if "image" in rel.target_ref:
            image = rel.target_part
            image_filename = os.path.join(image_storage_path, image.partname.split("/")[-1])
            with open(image_filename, "wb") as img_file:
                img_file.write(image.blob)
    
    # Combine all text and table data into one string
    document_text = "\n".join(full_text)
    document_tables = "\n".join(tables_text)
    
    return document_text, document_tables


def load_doc_documents(file_path):
    # Load the document, extract text, tables, and save images
    text, tables = extract_text_and_images_and_tables(file_path)
    return text, tables

# def load_doc_documents(file_path):
#     document_loader = DoclingLoader(file_path)
#     docs = document_loader.load()
#     # text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
#     # chunks = text_splitter.split_documents(docs)
#     return docs


def summarize_with_llm(user_query, document_context, tables):
    """Retrieves relevant text using FAISS and summarizes it with an LLM."""
    prompt = f"""
    You are creating a short summary for a clinical trial. 
    - Your response should be **short and to the point** (maximum of 2 sentences).
    - It should be **engaging and easy to understand**.
    - **Avoid excessive medical jargon.**
    - **Highlight only the key points.**
    - If unsure, state that information is not available.
    
    **Question:** {user_query}  
    **Context:** {document_context}  
    **Infographic Summary:** 
    """
#     prompt = f"""
#     you are recruting participants for a clinical trial. 
#     - you should serve to highlight important points that may be deciding factor for potential participants who may or may not want to join.
#     - Finer details about doising, schedules etc should be limited.
#     - Always refer to the individual as a 'participant.'
#     - Use the provided context to answer the query in a Formal way. 
#     - Answer should be user readble and should be limited to maximum of 2 sentences.
#     - If unsure, state that you don't know.
#     - Optional details can be omitted if there isn't enough information.


# Query: {user_query}
# Context: {document_context} 
# Tables: {tables}
# Answer:
#     """
    # print(retrieved_text)
    response = LANGUAGE_MODEL.invoke(prompt)
    return response

raw_docs, tables = load_doc_documents('./HRP-503 - SAMPLE Biomedical Protocol.docx')
# ai_response=summarize_with_llm(user_query, raw_docs)
# print(ai_response)
# response_dict = dict()
for each_q in user_query:
    x = summarize_with_llm(each_q, raw_docs, tables)
    print(x)
# print(response_dict)

# df = pd.DataFrame(list(response_dict.items()), columns=['Query', 'Answer'])
# # pd.ExcelWriter(df)
# df.to_excel('qa_2.xlsx', index=False)
# ai_response = summarize_with_llm(raw_docs)
# print(ai_response)
